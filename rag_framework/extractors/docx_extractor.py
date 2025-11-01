"""Extracteur basé sur python-docx (extraction de documents Word)."""

from pathlib import Path
from typing import Any, ClassVar

from rag_framework.extractors.base import BaseExtractor, ExtractionResult
from rag_framework.utils.logger import get_logger

logger = get_logger(__name__)


class DocxExtractor(BaseExtractor):
    """Extracteur utilisant python-docx pour l'extraction de documents Word.

    python-docx est la librairie de référence en 2025 pour l'extraction
    de documents Microsoft Word (.docx et .doc via conversion).
    Plus rapide et précis que Docling pour les documents Word simples.

    Avantages:
    - Extraction native et rapide de .docx
    - Préservation des paragraphes et styles
    - Extraction de tableaux structurés
    - Support des en-têtes et pieds de page
    - Extraction de métadonnées (auteur, date, etc.)
    - Pas de dépendance lourde (contrairement à Docling)

    Limitations:
    - Pas de support direct des .doc (ancien format)
    - Images et graphiques non extraits (seulement texte)
    - Formules mathématiques complexes non supportées
    - Commentaires et révisions ignorés

    Parameters
    ----------
    config : dict[str, Any]
        Configuration de l'extracteur.
        Clés supportées:
        - extract_tables : bool (défaut: True)
        - extract_headers_footers : bool (défaut: True)
        - preserve_formatting : bool (défaut: False)
        - min_text_length : int (défaut: 10)
        - extract_metadata : bool (défaut: True)

    Notes:
    -----
    Pour les documents Word complexes avec images, OCR ou mise en page
    sophistiquée, Docling peut être utilisé en fallback.
    """

    SUPPORTED_EXTENSIONS: ClassVar[set[str]] = {".docx", ".docm"}

    def can_extract(self, file_path: Path) -> bool:
        """Vérifie si le fichier est un document Word moderne (.docx/.docm).

        Parameters
        ----------
        file_path : Path
            Chemin vers le fichier.

        Returns:
        -------
        bool
            True si le fichier a l'extension .docx ou .docm.

        Examples:
        --------
        >>> extractor = DocxExtractor(config={})
        >>> extractor.can_extract(Path("rapport.docx"))
        True
        >>> extractor.can_extract(Path("rapport.doc"))
        False
        """
        return file_path.suffix.lower() in self.SUPPORTED_EXTENSIONS

    def extract(self, file_path: Path) -> ExtractionResult:
        """Extrait le texte d'un document Word avec python-docx.

        Parameters
        ----------
        file_path : Path
            Chemin vers le fichier Word.

        Returns:
        -------
        ExtractionResult
            Résultat de l'extraction.

        Notes:
        -----
        L'extraction combine les paragraphes, tableaux, en-têtes et
        pieds de page pour reconstituer le document complet.
        """
        try:
            # Import tardif pour éviter erreur si librairie non installée
            from docx import Document

            # Options d'extraction
            extract_tables = self.config.get("extract_tables", True)
            extract_headers_footers = self.config.get("extract_headers_footers", True)
            preserve_formatting = self.config.get("preserve_formatting", False)

            # Ouverture du document
            doc = Document(str(file_path))

            text_parts = []

            # Extraction des en-têtes (si activé)
            if extract_headers_footers:
                for section in doc.sections:
                    header_text = self._extract_header_footer(section.header)
                    if header_text:
                        text_parts.append(f"=== En-tête ===\n{header_text}")

            # Extraction du corps du document
            paragraph_count = 0
            table_count = 0

            for element in doc.element.body:
                # Paragraphe
                if element.tag.endswith("p"):
                    # Trouver le paragraphe correspondant
                    for para in doc.paragraphs:
                        if para._element == element:
                            para_text = para.text.strip()
                            if para_text:
                                if preserve_formatting:
                                    # Préservation du style (gras, italique, etc.)
                                    formatted_text = self._format_paragraph(para)
                                    text_parts.append(formatted_text)
                                else:
                                    text_parts.append(para_text)
                                paragraph_count += 1
                            break

                # Tableau
                elif element.tag.endswith("tbl") and extract_tables:
                    # Trouver le tableau correspondant
                    for table in doc.tables:
                        if table._element == element:
                            table_text = self._extract_table(table)
                            if table_text:
                                table_count += 1
                                text_parts.append(
                                    f"\n=== Tableau {table_count} ===\n{table_text}"
                                )
                            break

            # Extraction des pieds de page (si activé)
            if extract_headers_footers:
                for section in doc.sections:
                    footer_text = self._extract_header_footer(section.footer)
                    if footer_text:
                        text_parts.append(f"=== Pied de page ===\n{footer_text}")

            # Concaténation
            full_text = "\n\n".join(text_parts)

            # Métadonnées
            metadata: dict[str, Any] = {
                "file_size": file_path.stat().st_size,
                "file_name": file_path.name,
                "extractor": "python-docx",
                "paragraphs_count": paragraph_count,
                "tables_count": table_count,
            }

            # Extraction des métadonnées du document
            if self.config.get("extract_metadata", True):
                core_props = doc.core_properties
                if core_props.title:
                    metadata["title"] = core_props.title
                if core_props.author:
                    metadata["author"] = core_props.author
                if core_props.subject:
                    metadata["subject"] = core_props.subject
                if core_props.keywords:
                    metadata["keywords"] = core_props.keywords
                if core_props.created:
                    metadata["creation_date"] = str(core_props.created)
                if core_props.modified:
                    metadata["modification_date"] = str(core_props.modified)

            # Vérification de la longueur minimale
            min_length = self.config.get("min_text_length", 10)
            if len(full_text.strip()) < min_length:
                return ExtractionResult(
                    text=full_text,
                    success=False,
                    extractor_name=self.name,
                    metadata=metadata,
                    error=f"Texte extrait trop court ({len(full_text)} < {min_length})",
                    confidence_score=0.1,
                )

            # Score de confiance
            # python-docx est très fiable pour les .docx natifs
            confidence = 0.95 if len(full_text) > 50 else 0.7

            logger.debug(
                f"python-docx: Extrait {len(full_text)} caractères "
                f"({paragraph_count} paragraphes, {table_count} tableaux) "
                f"(confidence={confidence:.2f})"
            )

            return ExtractionResult(
                text=full_text,
                success=True,
                extractor_name=self.name,
                metadata=metadata,
                confidence_score=confidence,
            )

        except ImportError:
            error_msg = "python-docx n'est pas installé. Installez avec: pip install python-docx"
            logger.error(error_msg)
            return ExtractionResult(
                text="",
                success=False,
                extractor_name=self.name,
                metadata={},
                error=error_msg,
                confidence_score=0.0,
            )

        except Exception as e:
            error_msg = f"Erreur python-docx extraction: {e}"
            logger.warning(error_msg)
            return ExtractionResult(
                text="",
                success=False,
                extractor_name=self.name,
                metadata={},
                error=error_msg,
                confidence_score=0.0,
            )

    def _extract_header_footer(self, header_or_footer: Any) -> str:
        """Extrait le texte d'un en-tête ou pied de page.

        Parameters
        ----------
        header_or_footer : HeaderPart or FooterPart
            Objet en-tête ou pied de page de python-docx.

        Returns:
        -------
        str
            Texte extrait.
        """
        paragraphs = []
        for para in header_or_footer.paragraphs:
            text = para.text.strip()
            if text:
                paragraphs.append(text)
        return "\n".join(paragraphs)

    def _extract_table(self, table: Any) -> str:
        """Extrait le contenu d'un tableau au format Markdown.

        Parameters
        ----------
        table : Table
            Objet tableau de python-docx.

        Returns:
        -------
        str
            Tableau formaté en Markdown.
        """
        rows = []
        for row_idx, row in enumerate(table.rows):
            cells = [cell.text.strip() for cell in row.cells]
            # Format Markdown
            rows.append("| " + " | ".join(cells) + " |")

            # Ligne de séparation après l'en-tête (première ligne)
            if row_idx == 0 and len(cells) > 0:
                separator = "|" + "|".join(["---"] * len(cells)) + "|"
                rows.append(separator)

        return "\n".join(rows)

    def _format_paragraph(self, para: Any) -> str:
        """Formate un paragraphe en préservant les styles (gras, italique).

        Parameters
        ----------
        para : Paragraph
            Objet paragraphe de python-docx.

        Returns:
        -------
        str
            Texte formaté avec marqueurs Markdown.

        Notes:
        -----
        Cette fonction est utilisée uniquement si preserve_formatting=True.
        """
        formatted_parts = []

        for run in para.runs:
            text = run.text
            if run.bold and run.italic:
                text = f"***{text}***"
            elif run.bold:
                text = f"**{text}**"
            elif run.italic:
                text = f"*{text}*"

            formatted_parts.append(text)

        return "".join(formatted_parts)
