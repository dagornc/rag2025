"""Adapter python-docx pour les fichiers Word (.docx).

Auteur: RAG Framework Team
Version: 1.0.0
"""

from typing import Any, ClassVar

from rag_framework.preprocessing.adapters.base import LibraryAdapter, ParsingError


class PythonDocxAdapter(LibraryAdapter):
    """Adapter pour la librairie python-docx.

    Extrait le texte des fichiers Word (.docx).

    Attributes:
        REQUIRED_MODULES: Liste des modules requis.
    """

    REQUIRED_MODULES: ClassVar[list[str]] = ["docx"]

    def parse(self, file_path: str) -> dict[str, Any]:
        """Parse un fichier Word avec python-docx.

        Args:
            file_path: Chemin vers le fichier .docx.

        Returns:
            Dictionnaire avec text, metadata, paragraphs.

        Raises:
            ParsingError: Si le parsing échoue.
        """
        try:
            from docx import Document

            # Ouvrir le document
            doc = Document(file_path)

            # Configuration
            extract_styles = self.library_config.get("extract_styles", True)

            # Extraire le texte de tous les paragraphes
            paragraphs = []
            full_text = []

            for para in doc.paragraphs:
                para_dict: dict[str, Any] = {
                    "text": para.text,
                    "index": len(paragraphs),
                }

                # Extraire les styles si configuré
                if extract_styles and para.style:
                    para_dict["style"] = para.style.name

                paragraphs.append(para_dict)
                if para.text.strip():
                    full_text.append(para.text)

            # Extraire le texte des tableaux
            tables_text = []
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        row_text.append(cell.text)
                    tables_text.append(" | ".join(row_text))

            if tables_text:
                full_text.extend(tables_text)

            # Métadonnées du document
            core_props = doc.core_properties
            metadata = {
                "title": core_props.title or "",
                "author": core_props.author or "",
                "subject": core_props.subject or "",
                "created": str(core_props.created) if core_props.created else "",
                "modified": str(core_props.modified) if core_props.modified else "",
                "paragraph_count": len(paragraphs),
                "table_count": len(doc.tables),
            }

            return {
                "text": "\n\n".join(full_text),
                "metadata": metadata,
                "paragraphs": paragraphs,
            }

        except Exception as e:
            raise ParsingError(f"Échec python-docx parsing : {e}") from e
